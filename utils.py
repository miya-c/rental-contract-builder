import os
import tempfile
import logging
from functools import wraps
from flask import render_template, flash, redirect, url_for
from flask_login import current_user
from jinja2 import Template
import weasyprint

from app import app, db
from models import Contract, Room, Building, Owner, RealEstateAgent, SpecialTerm, ContractTemplate

def require_admin(f):
    """Decorator to require admin role for a view"""
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if not current_user.is_authenticated or not current_user.is_admin():
            flash('この操作を実行する権限がありません。', 'danger')
            return redirect(url_for('dashboard'))
        return f(*args, **kwargs)
    return decorated_function

def get_contract_data(contract_id):
    """Get all data needed for the contract PDF generation"""
    contract = Contract.query.get(contract_id)
    if not contract:
        logging.error(f"Contract with ID {contract_id} not found")
        return None
    
    room = Room.query.get(contract.room_id)
    if not room:
        logging.error(f"Room with ID {contract.room_id} not found")
        return None
    
    building = Building.query.get(room.building_id)
    if not building:
        logging.error(f"Building with ID {room.building_id} not found")
        return None
    
    owner = Owner.query.get(building.owner_id)
    if not owner:
        logging.error(f"Owner with ID {building.owner_id} not found")
        return None
    
    agent = RealEstateAgent.query.get(contract.agent_id)
    if not agent:
        logging.error(f"Agent with ID {contract.agent_id} not found")
        return None
    
    # Get all special terms for this contract
    special_terms = contract.special_terms
    
    # Parse custom amenities JSON
    custom_amenities = []
    if room.custom_amenities:
        try:
            import json
            custom_amenities = json.loads(room.custom_amenities)
        except json.JSONDecodeError:
            logging.error(f"Error parsing custom amenities JSON: {room.custom_amenities}")
    
    # Get the contract template
    template = ContractTemplate.query.get(contract.template_id)
    if not template:
        logging.error(f"Template with ID {contract.template_id} not found")
        return None
    
    return {
        'contract': contract,
        'room': room,
        'building': building,
        'owner': owner,
        'agent': agent,
        'special_terms': special_terms,
        'custom_amenities': custom_amenities,
        'template': template
    }

def generate_pdf(contract_id):
    """Generate a PDF for the contract and return the file path"""
    data = get_contract_data(contract_id)
    if not data:
        return None
    
    # Create a temporary directory for PDF generation
    pdf_dir = os.path.join(tempfile.gettempdir(), 'lease_contracts')
    os.makedirs(pdf_dir, exist_ok=True)
    
    # Create a filename for the PDF
    pdf_filename = f"contract_{data['contract'].contract_number}.pdf"
    pdf_path = os.path.join(pdf_dir, pdf_filename)
    
    # Also create paths for original format files
    original_dir = os.path.join(pdf_dir, 'originals')
    os.makedirs(original_dir, exist_ok=True)
    
    try:
        template = data['template']
        contract = data['contract']
        
        # Create a more comprehensive replacement dictionary for all template types
        replace_dict = {
            '{{contract.contract_number}}': contract.contract_number,
            '{{contract.tenant_name}}': contract.tenant_name,
            '{{contract.tenant_address}}': contract.tenant_address,
            '{{contract.tenant_phone}}': contract.tenant_phone or '',
            '{{contract.tenant_email}}': contract.tenant_email or '',
            '{{contract.start_date}}': contract.start_date.strftime('%Y年%m月%d日') if contract.start_date else '',
            '{{contract.end_date}}': contract.end_date.strftime('%Y年%m月%d日') if contract.end_date else '',
            '{{contract.rent_amount}}': f"{contract.rent_amount:,}円",
            '{{contract.security_deposit}}': f"{contract.security_deposit:,}円" if contract.security_deposit else '-',
            '{{contract.key_money}}': f"{contract.key_money:,}円" if contract.key_money else '-',
            '{{contract.management_fee}}': f"{contract.management_fee:,}円" if contract.management_fee else '-',
            '{{owner.name}}': data['owner'].name,
            '{{owner.address}}': data['owner'].address,
            '{{owner.phone}}': data['owner'].phone or '',
            '{{owner.email}}': data['owner'].email or '',
            '{{building.name}}': data['building'].name,
            '{{building.address}}': data['building'].address,
            '{{building.structure}}': data['building'].structure,
            '{{building.type}}': data['building'].building_type,
            '{{building.floors}}': str(data['building'].floors),
            '{{room.room_number}}': data['room'].room_number,
            '{{room.layout}}': data['room'].layout,
            '{{room.floor}}': str(data['room'].floor),
            '{{room.floor_area}}': f"{data['room'].floor_area}㎡",
            '{{agent.name}}': data['agent'].name,
            '{{agent.license_number}}': data['agent'].license_number,
        }
        
        # Handle different template types
        if template.file_type == 'html':
            # HTML template - Use Jinja2 to render the HTML template
            template_content = template.file_content
            jinja_template = Template(template_content)
            html_content = jinja_template.render(
                contract=data['contract'],
                room=data['room'],
                building=data['building'],
                owner=data['owner'],
                agent=data['agent'],
                special_terms=data['special_terms'],
                custom_amenities=data['custom_amenities']
            )
            
            # Save HTML file for reference
            html_path = os.path.join(original_dir, f"contract_{contract.contract_number}.html")
            with open(html_path, 'w', encoding='utf-8') as f:
                f.write(html_content)
            
            # Generate PDF with WeasyPrint
            weasyprint.HTML(string=html_content).write_pdf(pdf_path)
            
        elif template.file_type == 'excel':
            # Excel template processing
            from io import BytesIO
            import openpyxl
            
            # Load the Excel template
            if not template.file_binary:
                logging.error("Excel template has no binary data")
                return None
            
            workbook = openpyxl.load_workbook(BytesIO(template.file_binary))
            sheet = workbook.active
            
            # Replace placeholders in all cells
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                for row in sheet.rows:
                    for cell in row:
                        if cell.value and isinstance(cell.value, str):
                            for placeholder, value in replace_dict.items():
                                if placeholder in cell.value:
                                    cell.value = cell.value.replace(placeholder, value)
            
            # Save to a temporary Excel file
            excel_path = os.path.join(original_dir, f"contract_{contract.contract_number}.xlsx")
            workbook.save(excel_path)
            
            # For Excel to PDF conversion, we'll use WeasyPrint with an HTML table representation
            # First, convert Excel to HTML table
            html_content = "<html><body>"
            html_content += f"<h1>賃貸借契約書: {contract.contract_number}</h1>"
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                html_content += f"<h2>{sheet_name}</h2>"
                html_content += "<table border='1' style='border-collapse: collapse; width: 100%;'>"
                
                # Process all rows and cells
                for row in sheet.rows:
                    html_content += "<tr>"
                    for cell in row:
                        value = cell.value if cell.value is not None else ""
                        html_content += f"<td style='padding: 5px;'>{value}</td>"
                    html_content += "</tr>"
                html_content += "</table><br>"
            
            html_content += "</body></html>"
            
            # Save intermediate HTML
            html_path = os.path.join(original_dir, f"contract_{contract.contract_number}_excel.html")
            with open(html_path, 'w', encoding='utf-8') as f:
                f.write(html_content)
            
            # Generate PDF with WeasyPrint
            weasyprint.HTML(string=html_content).write_pdf(pdf_path)
            
            # Also save the Excel file as an attachment with the contract
            contract.original_file_path = excel_path
            
        elif template.file_type == 'word':
            # Word template processing
            from io import BytesIO
            import docx
            
            # Load the Word template
            if not template.file_binary:
                logging.error("Word template has no binary data")
                return None
                
            doc = docx.Document(BytesIO(template.file_binary))
            
            # Replace placeholders in paragraphs
            for para in doc.paragraphs:
                for placeholder, value in replace_dict.items():
                    if placeholder in para.text:
                        para.text = para.text.replace(placeholder, value)
            
            # Replace placeholders in tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            for placeholder, value in replace_dict.items():
                                if placeholder in para.text:
                                    para.text = para.text.replace(placeholder, value)
            
            # Save to a temporary Word file
            docx_path = os.path.join(original_dir, f"contract_{contract.contract_number}.docx")
            doc.save(docx_path)
            
            # For Word to PDF conversion, we'll use the same approach as Excel
            # Convert Word to HTML representation
            html_content = "<html><body>"
            html_content += f"<h1>賃貸借契約書: {contract.contract_number}</h1>"
            
            # Add each paragraph
            for para in doc.paragraphs:
                if para.text.strip():
                    html_content += f"<p>{para.text}</p>"
            
            # Add tables
            for table in doc.tables:
                html_content += "<table border='1' style='border-collapse: collapse; width: 100%;'>"
                for row in table.rows:
                    html_content += "<tr>"
                    for cell in row.cells:
                        cell_text = " ".join([p.text for p in cell.paragraphs])
                        html_content += f"<td style='padding: 5px;'>{cell_text}</td>"
                    html_content += "</tr>"
                html_content += "</table><br>"
            
            html_content += "</body></html>"
            
            # Save intermediate HTML
            html_path = os.path.join(original_dir, f"contract_{contract.contract_number}_word.html")
            with open(html_path, 'w', encoding='utf-8') as f:
                f.write(html_content)
            
            # Generate PDF with WeasyPrint
            weasyprint.HTML(string=html_content).write_pdf(pdf_path)
            
            # Also save the Word file as an attachment with the contract
            contract.original_file_path = docx_path
            
        elif template.file_type == 'pdf':
            # PDF template processing with PyPDF2
            from io import BytesIO
            import PyPDF2
            
            # Load the PDF template
            if not template.file_binary:
                logging.error("PDF template has no binary data")
                return None
            
            # For now, we'll create a copy of the PDF template
            # In a real implementation, you might use reportlab to add text overlays to the PDF
            pdf_template_path = os.path.join(original_dir, f"template_{template.id}.pdf")
            with open(pdf_template_path, 'wb') as f:
                f.write(template.file_binary)
            
            # Copy the template to the output PDF path
            with open(pdf_path, 'wb') as f:
                f.write(template.file_binary)
            
            # Store the source PDF path for reference
            contract.original_file_path = pdf_template_path
                
        else:
            logging.error(f"Unsupported template type: {template.file_type}")
            return None
            
        # If we got this far, the PDF was generated
        logging.info(f"Successfully generated PDF for contract {contract_id} at {pdf_path}")
        return pdf_path
        
    except Exception as e:
        logging.error(f"Error generating PDF: {str(e)}")
        import traceback
        logging.error(traceback.format_exc())
        return None
